#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Jul  5 16:34:42 2025

@author: abheekpathirana
"""

import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
from time import sleep
# === Dummy MCQ Bank ===


# === Predefined user database ===
USER_DB = {
    "user_1": "@1942_kreigsmarine_bismarch",
    "user_2": "@spf_1972_1354388",
    "user_lakshan": "1995@chaturanga_c_ika",
    "user_3": "@spf_1990_27872771",
    "user_4": "_1988_4436129",
    "user_5": "u@spf_1995_234467234",
    "user_6": "ni@spf_2000_38475",
    "user_7": "a@spf_1993_895678384",
    "user_8": "sachini@spf_204356701_10829",
    "user_9": "@spf_2001_17800829",
    "user_10": "_2001_10829",
    "user_11": "spf_2007891_10829",
    "user_admin": "admin_abheek",
    "user_12": "_1997_54928"
}

mcq_bank = {
    "Economics": [
        {
            "question": "What is the definition of opportunity cost?",
            "options": [
                "The money spent",
                "The next best alternative foregone",
                "Profit earned",
                "The total cost"
            ]
        },
        {
            "question": "Which of these is not a factor of production?",
            "options": [
                "Land",
                "Labour",
                "Technology",
                "Capital"
            ]
        },
        {
            "question": "Which of the following is not an economic good?",
            "options": [
                "Sunlight",
                "Tap Water",
                "Electricity",
                "Fried rice"
            ]
        },
        {
            "question": "What is meant by scarcity in economics?",
            "options": [
                "Lack of money",
                "Unlimited wants and limited resources",
                "Shortage of goods",
                "Recession in the economy"
            ]
        },
        {
            "question": "Which of the following is a primary industry?",
            "options": [
                "Mining",
                "Retailing",
                "Teaching",
                "Banking"
            ]
        },
        {
            "question": "Which sector is responsible for transforming raw materials into finished goods?",
            "options": [
                "Secondary sector",
                "Tertiary sector",
                "Primary sector",
                "Quaternary sector"
            ]
        },
        {
            "question": "What does GDP stand for?",
            "options": [
                "Gross Domestic Product",
                "General Domestic Pricing",
                "Government Development Plan",
                "Gross Developed Population"
            ]
        },
        {
            "question": "What is meant by specialization of labour?",
            "options": [
                "Concentration on a particular task",
                "Importing goods from abroad",
                "Increasing government spending",
                "Reducing working hours"
            ]
        },
        {
            "question": "Which statement best describes demand?",
            "options": [
                "The willingness and ability to buy a product at a given price",
                "The total supply of a product in the market",
                "The amount produced by firms",
                "The money spent by the government"
            ]
        },
            
        
    ],
    "ICT": [
        {
            "question": "What does CPU stand for?",
            "options": [
                "Central Processing Unit",
                "Control Power Unit",
                "Central Power Unit",
                "Computer Processing Unit"
            ]
        }
    ]
}

# === Simple login ===
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

def login_ui():
    st.title("Login to Springfield Exam Generator")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if username in USER_DB and USER_DB[username] == password:
            st.session_state.logged_in = True
            st.success("Login successful! Use the sidebar to continue.")
            sleep(0.1)
            st.success("A product by Abheek pathirana. ")
            sleep(0.1)
            st.success("AAP25™ " )
            sleep(0.1)
            st.success("double tap the login button.")
            
            
            
            st.stop()
        else:
            st.error("Invalid credentials.")

if not st.session_state.logged_in:
    login_ui()
    st.stop()

# === Page config ===
st.set_page_config(page_title="Springfield Exam Generator", layout="centered")

# === Sidebar navigation ===
page = st.sidebar.radio("Go to", ["Cover Page", "Question Structure", "Logout"])

# === Logout functionality ===
if page == "Logout":
    st.session_state.logged_in = False
    st.success("Logged out successfully.")
    st.stop

# === Session state init ===
if "cover_data" not in st.session_state:
    st.session_state.cover_data = {}
if "custom_mcqs" not in st.session_state:
    st.session_state.custom_mcqs = []

# === Page 1: Cover Page ===
if page == "Cover Page":
    st.title("Cover Page Metadata")

    subject = st.selectbox("Subject", list(mcq_bank.keys()))
    exam = st.selectbox("Term", ["First Term Examination-March", "Second Term Examination-July", "Third Term Examination-November"])
    teacher = st.text_input("Teacher Name")
    grade = st.selectbox("Grade", ["1", "2", "3", "4", "5", "6", "7", "8", "9", "9 Cambridge", "10", "10 Cambridge", "11"])
    duration = st.text_input("Time Duration (e.g., 1h 30min)")
    marks = st.text_input("Total Marks (e.g., 100)")

    if all([subject, exam, teacher, grade, duration, marks]):
        st.session_state.cover_data = {
            "subject": subject,
            "exam": exam,
            "teacher": teacher,
            "grade": grade,
            "duration": duration,
            "marks": marks
        }

        replacements = {
            "{{Subject}}": (subject, Pt(24)),
            "{{exam}}": (exam, Pt(24)),
            "{{teacher_n}}": (teacher, Pt(12)),
            "{{_Grade_}}": (grade, Pt(12)),
            "{{time_set_d}}": (duration, Pt(12)),
            "{{paper_marks_1}}": (marks, Pt(12)),
            "{{2022_d}}": ("2025", Pt(12))
        }

        def replace_placeholder(paragraph, replacements):
            for key, (val, size) in replacements.items():
                if key in paragraph.text:
                    text = ''.join(run.text for run in paragraph.runs)
                    if key in text:
                        text = text.replace(key, val)
                        for i in range(len(paragraph.runs) - 1, -1, -1):
                            paragraph._element.remove(paragraph.runs[i]._element)
                        new_run = paragraph.add_run(text)
                        new_run.font.size = size

        def fill_template(template_path, replacements):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                replace_placeholder(paragraph, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_placeholder(paragraph, replacements)
            return doc

        filled_doc = fill_template("spf_temp.docx", replacements)
        buffer = BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)

        st.success("Cover page ready!")
        st.download_button("Download Cover Page", data=buffer, file_name="cover_page.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("Fill in all fields to generate the cover page.")

# === Page 2: Question Structure ===
elif page == "Question Structure":
    st.title("MCQ Section Builder")

    if not st.session_state.cover_data:
        st.warning("Please complete the cover page first.")
    else:
        subject = st.session_state.cover_data["subject"]
        available_mcqs = mcq_bank.get(subject, [])
        custom_mcqs = st.session_state.custom_mcqs

        marks_per_mcq = st.number_input("Marks per MCQ?", min_value=1, max_value=10, step=1)
        options_mcq = st.selectbox("How many options per MCQ?", ["2", "3", "4", "5", "6"])
        options_mcq = int(options_mcq)

        st.subheader("Select Questions from MCQ Bank")
        selected_questions = []
        for i, mcq in enumerate(available_mcqs):
            with st.expander(f"{i+1}. {mcq['question']}"):
                for j, opt in enumerate(mcq["options"][:options_mcq]):
                    st.markdown(f"    {chr(65+j)}. {opt}")
                if st.checkbox("Include this question", key=f"mcq_bank_{i}"):
                    selected_questions.append(mcq)

        st.subheader("Add Custom MCQ")

        new_q_key = f"custom_q_{len(custom_mcqs)}"
        new_opts_keys = [f"custom_opt_{len(custom_mcqs)}_{i}" for i in range(options_mcq)]

        with st.form(key=new_q_key):
            custom_q = st.text_input("Enter custom question", key=new_q_key)
            custom_opts = [st.text_input(f"Option {chr(65+i)}", key=new_opts_keys[i]) for i in range(options_mcq)]

            submitted = st.form_submit_button("Add Custom MCQ")
            if submitted and custom_q and all(custom_opts):
                st.session_state.custom_mcqs.append({
                    "question": custom_q,
                    "options": custom_opts
                })
                st.success("Custom MCQ added!")

        st.subheader("Current Custom MCQs")
        for i, mcq in enumerate(custom_mcqs):
            st.markdown(f"**{i+1}. {mcq['question']}**")
            for j, opt in enumerate(mcq["options"]):
                st.markdown(f"    {chr(65+j)}. {opt}")

        if st.button("Generate Final Paper"):
            replacements = {
                "{{Subject}}": (subject, Pt(24)),
                "{{exam}}": (st.session_state.cover_data["exam"], Pt(24)),
                "{{teacher_n}}": (st.session_state.cover_data["teacher"], Pt(12)),
                "{{_Grade_}}": (st.session_state.cover_data["grade"], Pt(12)),
                "{{time_set_d}}": (st.session_state.cover_data["duration"], Pt(12)),
                "{{paper_marks_1}}": (st.session_state.cover_data["marks"], Pt(12)),
                "{{2022_d}}": ("2025", Pt(12))
            }

            def replace_placeholder(paragraph, replacements):
                for key, (val, size) in replacements.items():
                    if key in paragraph.text:
                        text = ''.join(run.text for run in paragraph.runs)
                        if key in text:
                            text = text.replace(key, val)
                            for i in range(len(paragraph.runs)-1, -1, -1):
                                paragraph._element.remove(paragraph.runs[i]._element)
                            new_run = paragraph.add_run(text)
                            new_run.font.size = size

            def fill_template(template_path, replacements):
                doc = Document(template_path)
                for paragraph in doc.paragraphs:
                    replace_placeholder(paragraph, replacements)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholder(paragraph, replacements)
                return doc

            doc = fill_template("spf_temp.docx", replacements)
            doc.add_page_break()
            doc.add_heading("Section A – Multiple Choice Questions", level=1)

            all_mcqs = selected_questions + custom_mcqs

            for i, mcq in enumerate(all_mcqs, 1):
                doc.add_paragraph(f"{i}. {mcq['question']} ({marks_per_mcq} marks)", style="Normal")
                for j, opt in enumerate(mcq["options"][:options_mcq]):
                    doc.add_paragraph(f"    {chr(65+j)}. {opt}")
                doc.add_paragraph("")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.success("Final paper ready!")
            st.download_button("Download Final Paper", data=buffer, file_name="final_exam_paper.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
